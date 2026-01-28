# app.py - Flask Backend
from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
import os
import shutil
import tempfile
import time
import json
import gc
from dotenv import load_dotenv
from werkzeug.utils import secure_filename

# LangChain imports
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_community.embeddings import HuggingFaceEmbeddings

# Docling imports 
try:
    from docling.document_converter import DocumentConverter
    
    try:
        from docling.document_converter import DocumentConverterConfig
        HAS_DOCLING_CONFIG = True
    except ImportError:
        try:
            from docling import DocumentConverterConfig
            HAS_DOCLING_CONFIG = True
        except ImportError:
            HAS_DOCLING_CONFIG = False
            print("Warning: DocumentConverterConfig not found, using default settings")
except ImportError as e:
    print(f"Error importing docling: {e}")
    HAS_DOCLING = False
else:
    HAS_DOCLING = True

# LangChain IBM integration
from langchain_ibm import WatsonxLLM

# Load environment variables
load_dotenv()

# Configuration from .env
IBM_URL = os.getenv("IBM_URL", "https://us-south.ml.cloud.ibm.com")
IBM_API_KEY = os.getenv("IBM_API_KEY")
IBM_PROJECT_ID = os.getenv("IBM_PROJECT_ID")

# Flask app setup
app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024  # 200MB max file size

# Create uploads directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize components
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
persist_directory = ".chromadb"

# Initialize Docling converter
def get_docling_converter():
    """Create and configure Docling document converter"""
    if not HAS_DOCLING:
        raise ImportError("Docling not installed")
    
    try:
        if HAS_DOCLING_CONFIG:
            
            config = DocumentConverterConfig()
            
            config.do_table_structure = True
            config.do_caption = True
            config.do_figure_boundary = True
            converter = DocumentConverter(config=config)
        else:
            
            converter = DocumentConverter()
            
            if hasattr(converter, 'do_table_structure'):
                converter.do_table_structure = True
            if hasattr(converter, 'do_caption'):
                converter.do_caption = True
            if hasattr(converter, 'do_figure_boundary'):
                converter.do_figure_boundary = True
        
        return converter
    except Exception as e:
        print(f"Error creating Docling converter: {e}")
        
        return DocumentConverter()

# Allowed file extensions
ALLOWED_EXTENSIONS = {
    'pdf', 'txt', 'docx', 'pptx', 
    'html', 'htm', 'md', 'rtf',
    'xlsx', 'xls', 'csv', 'tsv'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def check_ibm_credentials():
    """Verify IBM credentials are available"""
    if not IBM_API_KEY or not IBM_PROJECT_ID:
        return False, "IBM credentials not found. Please add IBM_API_KEY and IBM_PROJECT_ID to your .env file"
    return True, ""

# Function to clear database
def clear_database():
    """Delete the existing vector database with proper cleanup"""
    try:
        if os.path.exists(persist_directory):
            # Force close any open vectorstore connections
            try:
                gc.collect()
            except:
                pass
            
            # Add delay to ensure all file handles are released
            time.sleep(1)
            
            # Strategy 1: Try normal shutil.rmtree first
            try:
                shutil.rmtree(persist_directory)
                return True
            except (PermissionError, OSError):
                pass
            
            # Strategy 2: If normal deletion fails, try removing files recursively
            try:
                def remove_tree(path):
                    """Recursively remove all files and directories"""
                    if os.path.isdir(path):
                        for item in os.listdir(path):
                            item_path = os.path.join(path, item)
                            if os.path.isdir(item_path):
                                remove_tree(item_path)
                            else:
                                try:
                                    os.remove(item_path)
                                except:
                                    import stat
                                    os.chmod(item_path, stat.S_IWRITE)
                                    os.remove(item_path)
                        os.rmdir(path)
                
                remove_tree(persist_directory)
                return True
            except Exception as e:
                print(f"Error clearing database: {e}")
                return False
    except Exception as e:
        print(f"Error clearing database: {e}")
    return False

def extract_text_from_docling_result(doc_result, original_filename):
    """Extract text from Docling document result"""
    documents = []
    
    try:
        from langchain.schema import Document
        
        # 尝试获取文本内容
        text_content = ""
        
        # 方法1: 尝试使用export_to_text
        if hasattr(doc_result, 'export_to_text'):
            try:
                text_content = doc_result.export_to_text() or ""
            except:
                pass
        
        # 方法2: 尝试使用document属性
        if not text_content and hasattr(doc_result, 'document'):
            try:
                if hasattr(doc_result.document, 'export_to_text'):
                    text_content = doc_result.document.export_to_text() or ""
                elif hasattr(doc_result.document, '__str__'):
                    text_content = str(doc_result.document)
            except:
                pass
        
        # 方法3: 直接转换为字符串
        if not text_content:
            text_content = str(doc_result) if hasattr(doc_result, '__str__') else ""
        
        # 如果获取到了内容，创建文档
        if text_content.strip():
            doc = Document(
                page_content=text_content,
                metadata={
                    "source": original_filename,
                    "original_filename": original_filename,
                    "format": "processed_with_docling"
                }
            )
            documents.append(doc)
            print(f"Extracted {len(text_content)} characters from {original_filename}")
        
    except Exception as e:
        print(f"Error extracting text from docling result: {str(e)}")
    
    return documents

def load_document_with_docling(file_path, original_filename):
    """Load document using Docling"""
    try:
        if not HAS_DOCLING:
            raise ImportError("Docling not available")
        
        print(f"Processing {original_filename} with Docling...")
        
        # 初始化Docling转换器
        converter = get_docling_converter()
        
        # 转换文档
        result = converter.convert(file_path)
        
        print(f"Successfully converted {original_filename}")
        
        # 提取文本
        documents = extract_text_from_docling_result(result, original_filename)
        
        return documents
        
    except Exception as e:
        print(f"Error processing {original_filename} with Docling: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # 回退到基本文本提取
        return load_document_fallback(file_path, original_filename)

def load_document_fallback(file_path, original_filename):
    """Fallback document loading when Docling fails"""
    from langchain.schema import Document
    documents = []
    
    try:
        if original_filename.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={
                            "source": original_filename,
                            "original_filename": original_filename,
                            "format": "fallback_txt"
                        }
                    )
                    documents.append(doc)
                    print(f"Fallback: Extracted {len(content)} characters from TXT file")
                    
        elif original_filename.lower().endswith('.pdf'):
            # 使用PyPDF2作为PDF回退
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                content = ""
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        content += f"\n\n--- Page {page_num + 1} ---\n\n"
                        content += page_text
                
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={
                            "source": original_filename,
                            "original_filename": original_filename,
                            "format": "fallback_pdf",
                            "total_pages": len(reader.pages)
                        }
                    )
                    documents.append(doc)
                    print(f"Fallback: Extracted {len(content)} characters from PDF file")
            except ImportError:
                print("PyPDF2 not installed for PDF fallback")
                
        elif original_filename.lower().endswith('.docx'):
            # 使用python-docx作为DOCX回退
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                if content.strip():
                    doc = Document(
                        page_content=content,
                        metadata={
                            "source": original_filename,
                            "original_filename": original_filename,
                            "format": "fallback_docx"
                        }
                    )
                    documents.append(doc)
                    print(f"Fallback: Extracted {len(content)} characters from DOCX file")
            except ImportError:
                print("python-docx not installed for DOCX fallback")
                
    except Exception as e:
        print(f"Fallback extraction failed: {str(e)}")
    
    return documents

# Function to process documents
def process_documents(file_paths):
    """Process uploaded documents"""
    all_docs = []
    
    try:
        for file_path, original_filename in file_paths:
            print(f"\n--- Processing {original_filename} ---")
            
            # 首先尝试Docling
            documents = []
            if HAS_DOCLING:
                documents = load_document_with_docling(file_path, original_filename)
            
            # 如果Docling失败或不可用，使用回退
            if not documents:
                print(f"Using fallback for {original_filename}")
                documents = load_document_fallback(file_path, original_filename)
            
            if documents:
                print(f"Successfully extracted {len(documents)} document chunks from {original_filename}")
                all_docs.extend(documents)
            else:
                print(f"Warning: No content extracted from {original_filename}")
        
    except Exception as e:
        print(f"Error processing documents: {str(e)}")
        import traceback
        traceback.print_exc()
    
    return all_docs

# Routes
@app.route('/')
def index():
    """Serve the main HTML page"""
    return render_template('webpage.html')

@app.route('/api/check-database', methods=['GET'])
def check_database():
    """Check if database exists"""
    db_exists = os.path.exists(persist_directory)
    return jsonify({
        'database_exists': db_exists,
        'status': 'success'
    })

@app.route('/api/clear-database', methods=['POST'])
def api_clear_database():
    """Clear the vector database"""
    try:
        success = clear_database()
        if success:
            return jsonify({
                'status': 'success',
                'message': 'Database cleared successfully'
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Failed to clear database'
            }), 500
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/api/upload-documents', methods=['POST'])
def upload_documents():
    """Handle document upload and processing"""
    # Check credentials
    creds_ok, creds_msg = check_ibm_credentials()
    if not creds_ok:
        return jsonify({
            'status': 'error',
            'message': creds_msg
        }), 400
    
    # Check if files were uploaded
    if 'files' not in request.files:
        return jsonify({
            'status': 'error',
            'message': 'No files uploaded'
        }), 400
    
    files = request.files.getlist('files')
    if len(files) == 0 or files[0].filename == '':
        return jsonify({
            'status': 'error',
            'message': 'No files selected'
        }), 400
    
    # Get processing parameters
    chunk_size = request.form.get('chunk_size', 1000, type=int)
    chunk_overlap = request.form.get('chunk_overlap', 200, type=int)
    
    # Clear existing database
    if os.path.exists(persist_directory):
        clear_database()
    
    file_paths = []
    processed_files = []
    
    try:
        # Save uploaded files
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                file_paths.append((file_path, filename))
                processed_files.append(filename)
        
        if not file_paths:
            return jsonify({
                'status': 'error',
                'message': 'No valid files uploaded'
            }), 400
        
        print(f"\nStarting document processing for {len(file_paths)} files...")
        print(f"Docling available: {HAS_DOCLING}")
        
        # Process documents
        documents = process_documents(file_paths)
        
        if not documents:
            return jsonify({
                'status': 'error',
                'message': 'Failed to process documents. No content could be extracted.'
            }), 500
        
        print(f"\nTotal document sections extracted: {len(documents)}")
        
        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        splits = text_splitter.split_documents(documents)
        
        print(f"Created {len(splits)} chunks from {len(documents)} document sections")
        
        # Create and persist vector database
        print("Creating vector database...")
        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            persist_directory=persist_directory
        )
        vectorstore.persist()
        print("Vector database created and persisted successfully")
        
        # Clean up uploaded files
        for file_path, _ in file_paths:
            if os.path.exists(file_path):
                os.remove(file_path)
        
        return jsonify({
            'status': 'success',
            'message': f'Successfully processed {len(processed_files)} documents',
            'data': {
                'files_processed': len(processed_files),
                'total_chunks': len(splits),
                'total_sections': len(documents),
                'avg_chunk_size': sum(len(d.page_content) for d in splits)//len(splits) if splits else 0,
                'used_docling': HAS_DOCLING
            }
        })
        
    except Exception as e:
        print(f"Error in upload_documents: {str(e)}")
        import traceback
        traceback.print_exc()
        
        # Clean up any remaining files
        for file_path, _ in file_paths:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
        
        return jsonify({
            'status': 'error',
            'message': f'Error processing documents: {str(e)}'
        }), 500

@app.route('/api/ask-question', methods=['POST'])
def ask_question():
    """Handle question answering"""
    # Check credentials
    creds_ok, creds_msg = check_ibm_credentials()
    if not creds_ok:
        return jsonify({
            'status': 'error',
            'message': creds_msg
        }), 400
    
    # Check if database exists
    if not os.path.exists(persist_directory):
        return jsonify({
            'status': 'error',
            'message': 'Please upload and process documents first'
        }), 400
    
    # Get request data
    data = request.get_json()
    if not data or 'question' not in data:
        return jsonify({
            'status': 'error',
            'message': 'No question provided'
        }), 400
    
    question = data['question'].strip()
    k_value = data.get('k_value', 4)
    temperature = data.get('temperature', 0.1)
    
    if not question:
        return jsonify({
            'status': 'error',
            'message': 'Please enter a question'
        }), 400
    
    try:
        # Load existing vectorstore
        vectorstore = Chroma(
            persist_directory=persist_directory,
            embedding_function=embeddings
        )
        
        # Initialize IBM Granite
        llm = WatsonxLLM(
            model_id="ibm/granite-3-3-8b-instruct",
            url=IBM_URL,
            apikey=IBM_API_KEY,
            project_id=IBM_PROJECT_ID,
            params={
                "temperature": temperature,
                "max_new_tokens": 512,
                "repetition_penalty": 1.1,
                "top_p": 0.9
            }
        )
        
        # Create prompt template
        prompt_template = """You are an intelligent document assistant. Use the following context from uploaded documents to answer the question. 
        If the answer cannot be found in the context, say "I cannot find this information in the provided documents."
        Provide clear citations to the source documents when possible.

        Context Information:
        {context}

        Question: {question}

        Please provide a detailed answer based on the context:"""
        
        PROMPT = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )
        
        # Setup RAG chain
        qa_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(search_kwargs={"k": k_value}),
            chain_type_kwargs={"prompt": PROMPT},
            return_source_documents=True
        )
        
        # Get answer
        start_time = time.time()
        result = qa_chain({"query": question})
        response_time = time.time() - start_time
        
        # Format sources
        sources = []
        for i, doc in enumerate(result["source_documents"], 1):
            sources.append({
                'id': i,
                'source': doc.metadata.get('source', 'Unknown Document'),
                'format': doc.metadata.get('format', 'unknown'),
                'content': doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content
            })
        
        return jsonify({
            'status': 'success',
            'data': {
                'answer': result["result"],
                'sources': sources,
                'metrics': {
                    'response_time': round(response_time, 2),
                    'sources_used': len(result["source_documents"]),
                    'confidence': min(len(result["source_documents"]) / k_value * 100, 100)
                }
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error processing question: {str(e)}'
        }), 500

if __name__ == '__main__':
    # Create templates directory if it doesn't exist
    os.makedirs('templates', exist_ok=True)
    
    # Move webpage.html to templates directory
    if os.path.exists('webpage.html'):
        shutil.move('webpage.html', 'templates/webpage.html')
    
    app.run(debug=True, port=5000)