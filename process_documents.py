# process_documents.py
import os
from typing import List
from pypdf import PdfReader
import docx
from PIL import Image
import pytesseract
from langchain_core.documents import Document
from pdf2image import convert_from_path

def extract_text_from_pdf(path: str) -> str:
    text = []
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text.append(content)
    except Exception:
        # fallback OCR if extraction fails
        try:
            pages = convert_from_path(path)
            for page in pages:
                text.append(pytesseract.image_to_string(page))
        except Exception as e:
            print(f"❌ Failed to extract PDF text from {path}: {e}")
    return "\n".join(text)

def extract_text_from_docx(path: str) -> str:
    try:
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:
        print(f"❌ Failed to extract DOCX text from {path}: {e}")
        return ""

def extract_text_from_image(path: str) -> str:
    try:
        img = Image.open(path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"❌ Failed to extract text from image {path}: {e}")
        return ""

def process_documents(sources: List[str]):
    docs = []
    for source in sources:
        if not os.path.exists(source):
            print(f"⚠️ File not found: {source}")
            continue

        ext = os.path.splitext(source)[1].lower()
        text = ""

        if ext == ".pdf":
            text = extract_text_from_pdf(source)
        elif ext == ".docx":
            text = extract_text_from_docx(source)
        elif ext in [".jpg", ".jpeg", ".png"]:
            text = extract_text_from_image(source)
        else:
            print(f"⚠️ Unsupported file format: {ext}")
            continue

        if text.strip():
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            for i, p in enumerate(paragraphs):
                docs.append(Document(page_content=p, metadata={"source": source, "chunk": i}))

    print(f"✅ {len(docs)} document chunks created.")
    return docs
